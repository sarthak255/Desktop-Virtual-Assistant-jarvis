# src/document_creation.py
# pip install python-docx

import os
from docx import Document
from tkinter import Tk, Button, filedialog

# Function to create a Word document
def create_document(content, file_path):
    doc = Document()
    doc.add_heading("Document Title", 0)
    doc.add_paragraph(content)
    doc.save(file_path)

# Function to open a file dialog
def open_file_dialog():
    root = Tk()
    root.withdraw()
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word documents", "*.docx")])
    return file_path

# GUI for document creation
def document_creation_gui():
    def on_create_document():
        content = "This is a sample document."
        file_path = open_file_dialog()
        if file_path:
            create_document(content, file_path)

    root = Tk()
    root.title("Document Creation")

    create_button = Button(root, text="Create Document", command=on_create_document)
    create_button.pack()

    root.mainloop()

# Test the document creation GUI
if __name__ == "__main__":
    document_creation_gui()
